import pypdf
from docx import Document
from io import BytesIO
from pathlib import Path

class DocumentParser:
    """Парсер резюме (PDF, DOCX). Короткий и надёжный."""
    
    def __init__(self, max_size_mb: int = 10):
        self.max_bytes = max_size_mb * 1024 * 1024

    def parse(self, file_bytes: bytes, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        
        if len(file_bytes) > self.max_bytes:
            raise ValueError(f"Файл слишком большой (макс. {self.max_bytes // (1024*1024)} МБ)")
        if ext == ".doc":
            raise ValueError("Формат .doc устарел. Сохраните файл как .docx")
        if ext == ".pdf":
            return self._parse_pdf(file_bytes)
        if ext == ".docx":
            return self._parse_docx(file_bytes)
            
        raise ValueError(f"Неподдерживаемый формат: {ext}. Используйте PDF или DOCX")

    def _parse_pdf(self, file_bytes: bytes) -> str:
        try:
            reader = pypdf.PdfReader(BytesIO(file_bytes))
            pages_text = []
            
            for page in reader.pages:
                if text := page.extract_text():
                    pages_text.append(text.strip())
            
            if not pages_text:
                raise ValueError("PDF не содержит текста (возможно, это скан-изображение).")
                
            return "\n\n".join(pages_text)
        except Exception as e:
            raise ValueError(f"Ошибка чтения PDF: {e}")

    def _parse_docx(self, file_bytes: bytes) -> str:
        try:
            doc = Document(BytesIO(file_bytes))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Извлекаем текст из таблиц (важно для резюме)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            
            if not parts:
                raise ValueError("DOCX файл пуст или не содержит текста.")
                
            return "\n".join(parts)
        except Exception as e:
            if "zip" in str(e).lower():
                raise ValueError("Файл повреждён. Откройте его в Word и пересохраните как .docx")
            raise ValueError(f"Ошибка чтения DOCX: {e}")